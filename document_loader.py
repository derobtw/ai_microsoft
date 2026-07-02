import io
import xml.etree.ElementTree as ET
from pathlib import Path
from zipfile import ZipFile

import fitz
import pdfplumber
from docx import Document
from PIL import Image


try:
    import pytesseract
except ImportError:
    pytesseract = None


SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}
CHUNK_SIZE = 550
CHUNK_OVERLAP = 60


def clean_text(text):
    return " ".join(str(text).split())


def try_ocr_image(image_bytes):
    if pytesseract is None:
        return ""

    try:
        image = Image.open(io.BytesIO(image_bytes))
        text = pytesseract.image_to_string(image)
        return clean_text(text)
    except Exception:
        return ""


def read_txt_file(file_path):
    text = Path(file_path).read_text(encoding="utf-8", errors="ignore")

    return [
        {
            "content_type": "text",
            "location": "text file",
            "text": text,
        }
    ]


def read_pdf_file(file_path):
    items = []

    with pdfplumber.open(str(file_path)) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""

            if text.strip():
                items.append(
                    {
                        "content_type": "pdf_text",
                        "location": f"page {page_number}",
                        "text": text,
                    }
                )

            tables = page.extract_tables()

            for table_index, table in enumerate(tables, start=1):
                table_lines = [f"PDF table {table_index} on page {page_number}:"]

                for row in table:
                    cleaned_cells = [clean_text(cell) if cell else "" for cell in row]
                    table_lines.append(" | ".join(cleaned_cells))

                table_text = "\n".join(table_lines)

                if table_text.strip():
                    items.append(
                        {
                            "content_type": "pdf_table",
                            "location": f"page {page_number}, table {table_index}",
                            "text": table_text,
                        }
                    )

    pdf_document = fitz.open(str(file_path))

    for page_index in range(len(pdf_document)):
        page = pdf_document[page_index]
        images = page.get_images(full=True)

        for image_index, image_info in enumerate(images, start=1):
            xref = image_info[0]
            image_data = pdf_document.extract_image(xref)
            image_bytes = image_data.get("image", b"")
            image_extension = image_data.get("ext", "unknown")

            ocr_text = try_ocr_image(image_bytes)

            if ocr_text:
                image_text = (
                    f"PDF image {image_index} on page {page_index + 1}. "
                    f"OCR text: {ocr_text}"
                )
            else:
                image_text = (
                    f"PDF image {image_index} on page {page_index + 1}. "
                    f"Image format: {image_extension}. "
                    "No readable OCR text was found."
                )

            items.append(
                {
                    "content_type": "pdf_image",
                    "location": f"page {page_index + 1}, image {image_index}",
                    "text": image_text,
                }
            )

    pdf_document.close()
    return items


def extract_docx_xml_text(file_path):
    texts = []
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    with ZipFile(file_path) as docx_zip:
        xml_files = [
            name
            for name in docx_zip.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        ]

        for xml_file in xml_files:
            xml_content = docx_zip.read(xml_file)
            root = ET.fromstring(xml_content)

            for text_node in root.findall(".//w:t", namespaces):
                if text_node.text and text_node.text.strip():
                    texts.append(text_node.text.strip())

    return texts


def read_docx_file(file_path):
    document = Document(str(file_path))
    items = []
    seen_texts = set()

    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        text = clean_text(paragraph.text)

        if text and text not in seen_texts:
            items.append(
                {
                    "content_type": "docx_paragraph",
                    "location": f"paragraph {paragraph_index}",
                    "text": text,
                }
            )
            seen_texts.add(text)

    for table_index, table in enumerate(document.tables, start=1):
        table_lines = [f"DOCX table {table_index}:"]

        for row in table.rows:
            cell_texts = []

            for cell in row.cells:
                cell_text = clean_text(
                    " ".join(
                        paragraph.text
                        for paragraph in cell.paragraphs
                        if paragraph.text.strip()
                    )
                )

                if cell_text and cell_text not in cell_texts:
                    cell_texts.append(cell_text)

            if cell_texts:
                table_lines.append(" | ".join(cell_texts))

        table_text = "\n".join(table_lines)

        if len(table_lines) > 1 and table_text not in seen_texts:
            items.append(
                {
                    "content_type": "docx_table",
                    "location": f"table {table_index}",
                    "text": table_text,
                }
            )
            seen_texts.add(table_text)

    

    image_count = 0

    for relationship in document.part.rels.values():
        if "image" not in relationship.reltype:
            continue

        image_count += 1
        image_part = relationship.target_part
        image_bytes = image_part.blob
        ocr_text = try_ocr_image(image_bytes)

        if ocr_text:
            image_text = f"DOCX image {image_count}. OCR text: {ocr_text}"
        else:
            image_text = f"DOCX image {image_count}. No readable OCR text was found."

        items.append(
            {
                "content_type": "docx_image",
                "location": f"image {image_count}",
                "text": image_text,
            }
        )

    return items


def read_document(file_path):
    file_path = Path(file_path)
    extension = file_path.suffix.lower()

    if extension == ".txt":
        return read_txt_file(file_path)

    if extension == ".pdf":
        return read_pdf_file(file_path)

    if extension == ".docx":
        return read_docx_file(file_path)

    raise ValueError(f"Unsupported file type: {file_path.name}")


def chunk_text(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    chunks = []
    text = clean_text(text)

    if not text:
        return chunks

    start = 0

    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end].strip()

        if chunk:
            chunks.append(chunk)

        start += chunk_size - overlap

    return chunks
